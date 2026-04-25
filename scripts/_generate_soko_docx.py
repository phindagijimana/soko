"""One-off generator for soko.docx — run from repo root: python3 scripts/_generate_soko_docx.py"""
from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_para(document: Document, text: str, bold: Optional[str] = None) -> None:
    p = document.add_paragraph()
    if bold:
        r = p.add_run(bold)
        r.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def main() -> None:
    root = Path(__file__).resolve().parent.parent
    out = root / "soko.docx"

    doc = Document()
    t = doc.add_heading("Soko — requirements for developers and use cases", 0)
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p = doc.add_paragraph()
    p.add_run("Version: 1.0 (aligned to current Soko / Agri Marketplace codebase)").italic = True
    p = doc.add_paragraph()
    p.add_run("Audience: software developers, DevOps, and product owners building or operating the pilot.").italic = True

    add_heading(doc, "1. Purpose of this document", 1)
    doc.add_paragraph(
        "This document captures functional expectations (use cases) and non-functional / engineering "
        "requirements so new developers can run, extend, and deploy Soko with confidence. It reflects the "
        "pilot design for Rwanda and East Africa: phone-based profiles, listings, orders, trust, and admin support."
    )

    add_heading(doc, "2. Product summary", 1)
    doc.add_paragraph(
        "Soko is a web-based agricultural marketplace. Guests may browse public listings and reviews. "
        "Registered users (farmers and buyers) authenticate with phone and OTP, then perform role-specific "
        "actions. Administrators manage verification, support, and read operational metrics. In-app card "
        "or mobile payments are not in scope for this pilot; settlement is expected outside the application."
    )

    add_heading(doc, "3. Actors", 1)
    add_bullets(
        doc,
        [
            "Guest: unauthenticated visitor; can browse and read public data.",
            "Farmer: authenticated user with role farmer; can create listings, manage incoming orders, request verification.",
            "Buyer: authenticated user with role buyer; can place orders and submit reviews (subject to business rules).",
            "Admin: user flagged is_admin; can access Soko Admin (metrics, verification queue, audit, support tickets).",
        ],
    )

    add_heading(doc, "4. Developer requirements", 1)

    add_heading(doc, "4.1 Technology stack", 2)
    add_bullets(
        doc,
        [
            "Backend: Python 3.12, FastAPI, Uvicorn, SQLAlchemy 2, Alembic, Pydantic v2, httpx, pytest.",
            "Auth: phone + OTP, bearer tokens; Twilio (or file/console) for SMS in development.",
            "Database: SQLite for local; PostgreSQL recommended for production with Alembic migrations.",
            "Frontend: React 18, Vite 5, react-router-dom v6, static build for production.",
            "Deployment: static UI can be served from GitHub Pages; API must be hosted on a public HTTPS origin.",
        ],
    )

    add_heading(doc, "4.2 Repository layout (expectation)", 2)
    add_bullets(
        doc,
        [
            "backend/ — FastAPI app (app/main.py, models, settings, storage, tests).",
            "frontend/ — Vite + React; public/ for favicon, .nojekyll; dist/ is build output (not committed).",
            "infra/ — e.g. Nginx when using Docker Compose.",
            "docker-compose.yml — optional Postgres, API, Nginx (and S3 profile).",
            "./soko — project CLI: install, start, stop, status, restart (local dev).",
        ],
    )

    add_heading(doc, "4.3 Local development", 2)
    add_bullets(
        doc,
        [
            "Run ./soko install then ./soko start from the repository root (creates venv, installs deps, starts API + Vite).",
            "Backend alone: cd backend && uvicorn app.main:app --reload --port 2500 (after pip install -r requirements.txt).",
            "Frontend alone: cd frontend && npm run dev (expects API URL; Vite may use .env.development.local from ./soko start).",
            "Run backend tests: cd backend && pytest -q. Frontend: npm run build to verify the bundle.",
        ],
    )

    add_heading(doc, "4.4 Environment and configuration (production expectations)", 2)
    add_bullets(
        doc,
        [
            "Set ENVIRONMENT=production, DEBUG=false, and a strong SECRET_KEY; never use defaults in public deployments.",
            "Configure DATABASE_URL (PostgreSQL) and run alembic upgrade head; set CREATE_TABLES_ON_STARTUP=false when DB is managed only via migrations.",
            "ALLOWED_ORIGINS: must include the exact browser origin of the static site (e.g. https://<user>.github.io).",
            "TRUSTED_HOSTS: include the public hostname of the API so TrustedHostMiddleware allows requests.",
            "SMS: SMS_PROVIDER=twilio with real AC… Account SID and valid TWILIO_*; test delivery in a staging project first.",
            "Storage: for multi-instance or scale, set STORAGE_BACKEND=s3 with bucket, keys, and MEDIA_PUBLIC_BASE_URL (images must load in the browser).",
            "SENTRY_DSN: optional, for error monitoring; health endpoint reports Sentry state.",
            "TLS: terminate HTTPS at a load balancer or reverse proxy; repository ships HTTP-oriented examples only.",
            "CORS: production uses explicit origins (not *); see settings for dev vs production behavior.",
        ],
    )

    add_heading(doc, "4.5 Frontend build and hosting", 2)
    add_bullets(
        doc,
        [
            "Vite base path must match GitHub project Pages: basename import.meta.env.BASE_URL (e.g. /soko/).",
            "VITE_API_URL: set at build time to the public API base URL (no trailing slash); required for a working UI against a remote API.",
            "GitHub Pages: workflow builds frontend and deploys dist; include 404.html = index.html for SPA deep links (e.g. /soko/admin).",
            "Two UI surfaces: public marketplace at /; Soko Admin at /admin (admin users only, others redirected to /).",
        ],
    )

    add_heading(doc, "4.6 API surface (summary for integrators)", 2)
    doc.add_paragraph(
        "The REST API is JSON over HTTPS. Public reads include GET /listings, GET /reviews, GET /recommendations. "
        "Write operations require Authorization: Bearer <token> except where noted. See OpenAPI (FastAPI) at /docs when the server is running."
    )
    add_bullets(
        doc,
        [
            "Auth: POST /users, POST /auth/request-otp, POST /auth/verify-otp, POST /auth/logout, GET /me",
            "Listings & media: GET/POST /listings, POST /images/upload",
            "Orders: POST/GET /orders, PATCH /orders/{id}",
            "Reviews: GET/POST /reviews",
            "Verification: POST /verification/request; admin GET/PATCH on /admin/verification-requests",
            "Support: POST/GET /support-tickets, admin PATCH /admin/support-tickets/{id}",
            "Engagement: POST /interactions, GET /recommendations",
            "Admin: GET /metrics/summary, GET /admin/audit-logs, GET /admin/users",
            "Health: GET /health, GET /ready",
        ],
    )

    add_heading(doc, "4.7 Engineering standards", 2)
    add_bullets(
        doc,
        [
            "Run tests before merging critical backend changes; keep domain rules in small modules to preserve readability.",
            "Rate limiting: in-memory for single instance; plan Redis or edge limits if you scale out horizontally.",
            "Audit log exists for key actions; do not log secrets in plain text responses.",
        ],
    )

    add_heading(doc, "5. Use cases to cover (functional)", 1)
    add_para(doc, "Each use case should be testable in staging with realistic SMS and a shared database policy.")

    uces = [
        (
            "UC-01 — Guest browses the marketplace",
            "Guest",
            "API reachable; optional seed or existing listings",
            "Open site; use search/filters. View recommended and all listings; read public reviews. May record interactions when logged in; guests do not need to sign in to browse.",
            "User understands offerings and trust signals without an account.",
        ),
        (
            "UC-02 — User registers a profile",
            "Prospective farmer or buyer",
            "Valid phone not yet registered; SMS or dev placeholder for OTP",
            "Submit name, phone, role, location on Create profile. Receive confirmation and proceed to OTP flow.",
            "User record created; can request OTP to sign in.",
        ),
        (
            "UC-03 — User signs in with phone OTP",
            "Registered user",
            "OTP service configured; rate limits and lockout rules respected",
            "Request OTP, enter code within expiry. Session token stored client-side. Optional logout invalidates server-side when implemented.",
            "Authenticated session; /me returns role and profile fields.",
        ),
        (
            "UC-04 — Farmer creates and manages listings",
            "Farmer",
            "Signed in as farmer",
            "Create listing (crop, quantity, price, location, description, optional image URL and/or file upload). Listings appear in search and market views. Farmer cannot place buyer orders with farmer profile (role guard).",
            "Listing persisted; media stored per STORAGE_BACKEND; farmer can receive orders for own listings.",
        ),
        (
            "UC-05 — Buyer places and tracks an order",
            "Buyer",
            "Signed in as buyer; valid listing",
            "Submit order for listing and quantity. Order appears in buyer’s list. Farmer can update status (e.g. accepted, rejected, completed) per business rules. SMS notifications use configured provider in production.",
            "Order state visible to both parties; review flow can reference completed orders where applicable.",
        ),
        (
            "UC-06 — Buyer leaves a review",
            "Buyer",
            "Signed in as buyer; rules on order linkage as implemented",
            "Submit review for a farmer (score, text, optional order id). Review appears in public list and supports trust in recommendations.",
            "New review stored; public trust signals update.",
        ),
        (
            "UC-07 — User requests account verification",
            "Farmer (or as implemented)",
            "Signed in",
            "Submit verification request with document metadata. Request enters pending state for admin review.",
            "Admin can later approve or reject; user verification flag updates accordingly.",
        ),
        (
            "UC-08 — User opens a support ticket",
            "Signed-in user",
            "Valid session",
            "Create ticket with category (general, dispute, abuse, bug), subject, message. User sees their tickets; admin sees all in Soko Admin.",
            "Ticket stored with status; admin can add notes and change status (in_progress, resolved, closed).",
        ),
        (
            "UC-09 — Admin uses Soko Admin",
            "Admin (is_admin)",
            "Admin phone configured in ADMIN_PHONE_NUMBERS; user promoted or seeded as admin",
            "From public site, follow “Soko Admin” to /admin. View metrics, verification queue, audit log, and moderate all support tickets. Return to public site via back link. Non-admins cannot access /admin (redirect to /).",
            "Operational visibility and moderation without exposing admin UI to the general marketplace navigation tabs.",
        ),
        (
            "UC-10 — Listings and recommendations use engagement data",
            "Signed-in user",
            "Token present",
            "Track views/clicks; recommendations blend trust, recency, location, and behavior per implemented ranking.",
            "Personalized or ranked experiences improve for active users while guests still see default ranking.",
        ),
        (
            "UC-11 — System health and readiness",
            "Operator / DevOps",
            "Service deployed with dependencies (DB, optional S3, SMS)",
            "GET /health and /ready for probes; optional Sentry enabled based on SENTRY_DSN. Backups run per runbook (SQLite or pg_dump).",
            "Team can monitor uptime and data protection expectations.",
        ),
    ]
    for title, actor, pre, flow, post in uces:
        add_heading(doc, title, 2)
        doc.add_paragraph(
            f"Actor: {actor}. Preconditions: {pre}. Main flow: {flow} Success: {post}."
        )

    add_heading(doc, "6. Explicitly out of scope (current pilot)", 1)
    add_bullets(
        doc,
        [
            "In-app or embedded payment collection (e.g. card, MoMo API) — business settlement is out of band.",
            "Full legal KYC/ID document upload pipeline beyond placeholder references — may be added later with storage and policy.",
            "Multi-language UI and accessibility audit beyond best-effort (future improvement).",
        ],
    )

    add_heading(doc, "7. Traceability to implementation", 1)
    doc.add_paragraph(
        "Backend: backend/app/main.py, models.py, settings.py, auth, storage, SMS. "
        "Frontend: frontend/src/App.jsx (routes, marketplace vs Soko Admin), lib/api.js. "
        "Operations: README.md, .env.example, docker-compose.yml, .github/workflows/deploy-github-pages.yml."
    )

    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
